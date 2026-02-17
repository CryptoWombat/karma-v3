"""Convert PRD Markdown to Word document."""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color="F2F2F2"):
    """Set background color for table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_markdown_table(doc, lines):
    """Parse markdown table and add to document."""
    if not lines:
        return
    rows = []
    for line in lines:
        line = line.strip()
        if not line or line.startswith("#"):
            break
        if "|" in line:
            cells = [c.strip() for c in line.split("|") if c.strip()]
            is_separator = cells and all(re.match(r"^\-+$", c.replace(" ", "")) for c in cells)
            if cells and not is_separator:
                rows.append(cells)
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for i, row_cells in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_cells):
            if j < len(row.cells):
                row.cells[j].text = cell_text
                if i == 0:
                    for paragraph in row.cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    doc.add_paragraph()


def parse_markdown_to_docx(md_path: Path, docx_path: Path):
    """Convert markdown file to Word document."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split("\n")
    i = 0
    in_code_block = False
    code_lines = []
    table_lines = []
    in_table = False

    while i < len(lines):
        line = lines[i]
        original_line = line

        if line.strip().startswith("```"):
            if in_code_block:
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                p.style = "Normal"
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if "|" in line and (line.strip().startswith("|") or (i > 0 and in_table)):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table and table_lines:
                add_markdown_table(doc, table_lines)
                table_lines = []
            in_table = False

        stripped = line.strip()
        if not stripped:
            if table_lines:
                add_markdown_table(doc, table_lines)
                table_lines = []
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.", stripped):
            p = doc.add_paragraph(re.sub(r"^\d+\.\s*", "", stripped), style="List Number")
        else:
            doc.add_paragraph(stripped)
        i += 1

    if table_lines:
        add_markdown_table(doc, table_lines)

    doc.save(docx_path)
    print(f"Created: {docx_path}")


if __name__ == "__main__":
    base = Path(__file__).parent
    md_file = base / "PRD_KARMA_PLATFORM_REBUILD.md"
    docx_file = base / "PRD_KARMA_PLATFORM_REBUILD.docx"
    parse_markdown_to_docx(md_file, docx_file)
